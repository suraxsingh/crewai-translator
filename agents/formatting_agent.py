# agents/formatting_agent.py
import docx

def rebuild_doc(translated_paragraphs, output_file):
    """
    Rebuild a .docx file from translated paragraphs.
    """
    new_doc = docx.Document()
    for para in translated_paragraphs:
        if para["type"] == "paragraph":
            new_doc.add_paragraph(para["content"])
        else:
            new_doc.add_paragraph(para["content"])  # placeholder for tables, titles, etc.
    new_doc.save(output_file)
    print(f"✅ Rebuilt translated document saved at: {output_file}")
